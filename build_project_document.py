from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Ekmind_AI_Content_Studio_Project_Document.docx"

NAVY = "112B45"
DEEP_NAVY = "081C2D"
GOLD = "F2B705"
PALE_GOLD = "FFF4CC"
PALE_BLUE = "EAF1F7"
MID_BLUE = "446A85"
WHITE = "FFFFFF"
INK = "243442"
GREY = "667684"
GREEN = "18794E"
RED = "B42318"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=120, bottom=110, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=9, align=WD_ALIGN_PARAGRAPH.LEFT) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_col_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_table(table, headers: list[str], rows: list[list[str]], widths: list[float], *, compact=False) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], NAVY)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE, size=8.5 if compact else 9)
    set_repeat_table_header(table.rows[0])
    for r_idx, values in enumerate(rows, start=1):
        for c_idx, value in enumerate(values):
            cell = table.rows[r_idx].cells[c_idx]
            if r_idx % 2 == 0:
                set_cell_shading(cell, PALE_BLUE)
            align = WD_ALIGN_PARAGRAPH.CENTER if c_idx == 0 and len(headers) > 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cell, value, size=8 if compact else 8.7, align=align)
    set_col_widths(table, widths)


def add_table(doc, headers, rows, widths, *, compact=False):
    table = doc.add_table(rows=1, cols=len(headers))
    for row in rows:
        table.add_row()
    style_table(table, headers, rows, widths, compact=compact)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(GREY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.text = "EKMIND  /  AI CONTENT STUDIO"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.runs[0]
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MID_BLUE)
    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.65))
    table.autofit = False
    set_cell_text(table.cell(0, 0), "Project document | 4 September 2026", size=8, color=GREY)
    add_page_number(table.cell(0, 1).paragraphs[0])
    table.cell(0, 0).width = Inches(5.2)
    table.cell(0, 1).width = Inches(1.45)


def add_title(doc, title: str, subtitle: str | None = None, size: int | None = None) -> None:
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.page_break_before = True
    run = p.add_run(title)
    if size is not None:
        run.font.size = Pt(size)
    if subtitle:
        s = doc.add_paragraph(subtitle)
        s.style = doc.styles["Subtitle"]


def add_heading(doc, text: str, level=1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc, text: str, *, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        p.add_run(bold_lead).bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def add_bullets(doc, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(3)
        p.add_run(f"{index}.  {item}")


def add_callout(doc, title: str, body: str, fill=PALE_GOLD) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.12
    for run in p2.runs:
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(INK)


def page_break(doc) -> None:
    # Section titles carry page_break_before. Avoid standalone break paragraphs,
    # which can be pushed onto an otherwise blank page by a full preceding page.
    return None


def setup_styles(doc) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    add_header_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    title = doc.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(31)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_after = Pt(10)

    subtitle = doc.styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = RGBColor.from_string(MID_BLUE)

    for level, size in ((1, 20), (2, 14), (3, 11.5)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(NAVY if level < 3 else MID_BLUE)
        style.paragraph_format.space_before = Pt(10 if level == 1 else 7)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_after = Pt(3)


def build() -> None:
    doc = Document()
    setup_styles(doc)
    props = doc.core_properties
    props.title = "Ekmind AI Content Studio - Project Document"
    props.subject = "Architecture, workflow, evaluation, operations, and roadmap"
    props.author = "Ekmind"
    props.keywords = "AI, LangGraph, Streamlit, FastAPI, evaluation, HITL, content studio"

    # Cover
    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("EKMIND")
    r.font.name = "Aptos Display"
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(GOLD)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run("AI CONTENT\nSTUDIO")
    r.font.name = "Aptos Display"
    r.font.size = Pt(38)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph("PROJECT DOCUMENT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.runs[0]
    r.font.name = "Aptos"
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(MID_BLUE)
    doc.add_paragraph("\n")
    add_callout(doc, "Purpose", "A governed, human-in-the-loop system that turns an idea into researched thought leadership, branded slides, an avatar-led video, and a privately published content asset.")
    p = doc.add_paragraph("Version 1.1  |  4 September 2026")
    p.paragraph_format.space_before = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.color.rgb = RGBColor.from_string(GREY)
    p.runs[0].font.size = Pt(10)
    page_break(doc)

    # Document control and contents
    add_title(doc, "Document control", "A practical reference for stakeholders, developers, evaluators, and operators")
    add_table(doc, ["Field", "Value"], [
        ["Project", "Ekmind AI Content Studio"],
        ["Document version", "1.1"],
        ["Status", "Current implementation baseline"],
        ["Prepared", "4 September 2026"],
        ["Application location", r"Week4Assignment\ekmind-content-studio"],
        ["Primary interfaces", "Streamlit UI and FastAPI service"],
        ["Current local database", "SQLite; PostgreSQL can be configured later"],
    ], [1.7, 4.8])
    add_heading(doc, "Contents", 2)
    add_table(doc, ["Section", "Topic"], [
        ["1", "Executive summary"], ["2", "Objectives and scope"], ["3", "Solution architecture"],
        ["4", "RAG architecture and evidence workflow"], ["5", "End-to-end workflow"],
        ["6", "Core components and integrations"], ["7", "Research, grounding, and content quality"],
        ["8", "Human governance and controls"], ["9", "Evaluation framework"],
        ["10", "Current evaluation status"], ["11", "Setup, operation, and testing"],
        ["12", "Risks, limitations, and roadmap"],
        ["A", "API and configuration reference"], ["B", "Project artifacts"],
    ], [0.8, 5.7], compact=True)
    page_break(doc)

    # Executive summary
    add_title(doc, "1. Executive summary")
    add_body(doc, "Ekmind AI Content Studio is an agentic content-production application designed to convert a topic into a governed, reusable multimedia package. It combines planning, optional web research, three differentiated content angles, a LinkedIn-style post, automated evaluation, a branded PowerPoint deck, a narrated avatar video, captions, and optional private YouTube publication.")
    add_body(doc, "The system is intentionally not a one-shot text generator. LangGraph separates strategy and writing into stateful workflows, while deterministic application code controls persistence, approvals, retries, file generation, validation, provider calls, and publication safeguards.")
    add_heading(doc, "Current outcome", 2)
    add_table(doc, ["Capability", "Status", "Evidence"], [
        ["Core workflow", "Implemented", "Strategy, approval, post, slides, video, library"],
        ["Human-in-the-loop", "Implemented", "Mandatory approvals and reject/regenerate paths"],
        ["Evaluation suite", "Implemented", "24 governed metrics and versioned reports"],
        ["Automated tests", "Passing", "31 of 31 tests passed"],
        ["Validation split", "Passing", "8 of 8 deterministic workflow cases passed"],
        ["Release decision", "Not ready", "Faithfulness remains NOT_EVALUATED in the controlled report"],
        ["Live providers", "Configuration-dependent", "OpenAI, search, LangSmith, HeyGen, YouTube"],
    ], [2.0, 1.25, 3.25])
    add_callout(doc, "Evaluation position", "Overall evaluation readiness is approximately 7.5/10. Architecture, workflow controls, reproducibility, and trajectory testing are strong; human-calibrated faithfulness and real-provider validation remain the main evidence gaps.")
    add_heading(doc, "Design principle", 2)
    add_body(doc, "LLMs decide meaning; deterministic code executes operations. The model may interpret, plan, write, critique, and propose. It cannot bypass approval gates, execute arbitrary shell commands, choose unrestricted file paths, or silently publish content.")
    page_break(doc)

    # Objectives and scope
    add_title(doc, "2. Objectives and scope")
    add_heading(doc, "Business objectives", 2)
    add_bullets(doc, [
        "Accelerate the conversion of business ideas into credible thought-leadership content.",
        "Maintain human ownership of angles, factual claims, final copy, slides, video, and publication.",
        "Reuse approved content and capture evaluation evidence for continuous improvement.",
        "Create consistent branded assets across LinkedIn, presentation, and video formats.",
        "Make quality measurable through repeatable tests, trace evidence, and explicit release gates.",
    ])
    add_heading(doc, "In scope", 2)
    add_table(doc, ["Stage", "Delivered output"], [
        ["Strategy", "Research decision, evidence, and three differentiated angles"],
        ["Writing", "Draft, critique, bounded revisions, factual checks, approved post"],
        ["Presentation", "Eight-slide outline, branded PPTX, previews, speaker notes"],
        ["Video", "Scene plan, narration, avatar composition, captions, validated MP4"],
        ["Publishing", "Optional OAuth upload to YouTube with private visibility"],
        ["Memory", "Approved-content library and workspace-scoped persistence"],
        ["Evaluation", "Session metrics, regression suites, immutable reports, LangSmith export"],
    ], [1.55, 4.95])
    add_heading(doc, "Outside the current baseline", 2)
    add_bullets(doc, [
        "Direct LinkedIn publishing.",
        "Production PostgreSQL/pgvector operation in the current local environment.",
        "Proof that live-provider quality meets thresholds across a calibrated hidden-test set.",
        "A completed approved-post corpus large enough to validate Voice Match robustly.",
        "Public YouTube publication; the implemented upload path is private by design.",
    ])
    page_break(doc)

    # Architecture
    add_title(doc, "3. Solution architecture")
    add_body(doc, "The application uses a layered architecture so user interaction, orchestration, business rules, providers, persistence, and evaluation can evolve independently.")
    add_table(doc, ["Layer", "Technology", "Responsibility"], [
        ["Experience", "Streamlit", "Multi-page UI for creation, approvals, library, evaluation, and settings"],
        ["API", "FastAPI", "Session lifecycle, approvals, generation, export, and health endpoints"],
        ["Orchestration", "LangGraph", "Strategy and post graphs, routing, state, critic/revision loop"],
        ["AI", "OpenAI + Pydantic", "Structured planning, angles, writing, critique, and revisions"],
        ["Research", "Search provider", "Web retrieval, evidence metadata, relevance signals"],
        ["Persistence", "SQLAlchemy + SQLite", "Workspace-scoped session and approved-content records"],
        ["Presentation", "python-pptx + Pillow", "Template-based PPTX, previews, notes, and validation"],
        ["Video", "HeyGen + FFmpeg", "Avatar generation, slide composition, audio, captions, MP4 validation"],
        ["Observability", "LangSmith", "Tracing and approved-content dataset export when configured"],
        ["Evaluation", "Python + pytest", "Deterministic metrics, regression tests, reports, and hard gates"],
    ], [1.25, 1.65, 3.6], compact=True)
    add_heading(doc, "Logical flow", 2)
    add_callout(doc, "Interaction model", "Streamlit → FastAPI → ContentService → LangGraph / deterministic services → provider adapters → storage and evaluation. Session state is persisted after every meaningful transition, and the evaluation matrix is refreshed with the session.", PALE_BLUE)
    add_heading(doc, "Key architectural properties", 2)
    add_bullets(doc, [
        "Provider abstraction: mock and live implementations share controlled interfaces.",
        "Stateful workflow: generated artifacts and decisions remain associated with a session.",
        "Workspace boundaries: reads and library access are scoped by workspace identifiers.",
        "Recoverability: rejection records a reason and returns the workflow to the appropriate stage.",
        "Auditability: decision history, runtime metrics, evaluation rows, and optional traces are retained.",
    ])
    page_break(doc)

    # RAG architecture
    add_title(doc, "4. RAG architecture and evidence workflow")
    add_body(doc, "The current implementation uses retrieval-augmented generation over live web evidence. Research results are normalized into evidence records and supplied to the angle, writing, revision, and factual-evaluation stages. This is distinct from a production vector RAG platform: PostgreSQL/pgvector semantic retrieval and long-term embedding memory are planned capabilities, not part of the current local baseline.")
    add_heading(doc, "Retrieval-to-generation pipeline", 2)
    add_table(doc, ["Stage", "Processing", "Output / control"], [
        ["1. Route", "Planner interprets the topic and research mode", "Research required or skipped"],
        ["2. Retrieve", "Search provider executes the research query", "Ranked web results"],
        ["3. Normalize", "Capture passage, title, URL, publisher, rank, and relevance signal", "Structured Evidence records"],
        ["4. Context", "Format retrieved passages into bounded source context", "Evidence supplied to the LLM"],
        ["5. Generate", "Angles and post use evidence-only factual instructions", "Grounded draft with source links"],
        ["6. Extract", "Deterministic logic identifies factual and material claims", "Claim inventory"],
        ["7. Validate", "Compare each claim with focused evidence passages", "Support classification and score"],
        ["8. Govern", "Block material unsupported/contradicted claims; allow human review", "Approval or revision"],
    ], [1.05, 3.55, 2.05], compact=True)
    add_heading(doc, "Evidence object", 2)
    add_table(doc, ["Field", "Purpose"], [
        ["evidence_id", "Stable identifier for claim-to-source mapping"],
        ["claim / passage", "Retrieved text used as grounding context"],
        ["source_title and source_url", "Provenance and reviewer access"],
        ["publisher", "Source-origin signal"],
        ["retrieval_rank", "Original search position"],
        ["relevance_score / label", "Heuristic topical relevance signal"],
        ["verification_status", "Controlled or retrieved evidence state"],
    ], [2.05, 4.6], compact=True)
    add_title(doc, "4B. RAG controls and gaps")
    add_callout(doc, "RAG boundary", "The LLM can summarize and reason over retrieved evidence, but deterministic code owns retrieval records, workflow state, thresholds, approval gates, and artifact operations. When research is skipped, Evidence Coverage, Groundedness, Faithfulness, Unsupported Claims, and Factual Correctness are marked NOT_APPLICABLE with a Research only target.", PALE_BLUE)
    add_heading(doc, "Current gaps before production RAG", 2)
    add_bullets(doc, [
        "No production vector index, embedding pipeline, chunk lifecycle, or hybrid search.",
        "No automated source-authority ranking beyond the available provider and overlap signals.",
        "Long-term approved-content memory is stored, but semantic retrieval via pgvector is not active locally.",
        "Formal claim-evidence evaluator calibration is recommended for production, although not required for the course demonstration.",
    ])
    page_break(doc)

    # Workflow
    add_title(doc, "5. End-to-end workflow")
    add_table(doc, ["#", "Workflow stage", "System action", "Human control"], [
        ["1", "Idea intake", "Capture topic and research preference", "User initiates"],
        ["2", "Planning", "Interpret request and decide route", "Visible result"],
        ["3", "Research", "Retrieve and score evidence when required", "Sources reviewable"],
        ["4", "Angles", "Produce three differentiated strategies", "Approve or reject"],
        ["5", "Writing", "Draft, critique, and revise within limit", "Edit and approve/reject"],
        ["6", "Grounding", "Assess factual claims against evidence", "Review disputed claims"],
        ["7", "Library", "Save approved post and evaluation record", "Approval required"],
        ["8", "Slides", "Build outline, PPTX, previews, and notes", "Approve or reject"],
        ["9", "Video", "Plan scenes, generate avatar, compose captions", "Approve or reject"],
        ["10", "Publication", "Upload final video as private when configured", "Final approval"],
    ], [0.42, 1.25, 3.15, 1.65], compact=True)
    add_heading(doc, "LangGraph topology", 2)
    add_table(doc, ["Graph", "Nodes", "Termination rule"], [
        ["Strategy graph", "Understand request → retrieve research → generate angles", "Three candidate angles available"],
        ["Post graph", "Write → critique → revise → critique", "Quality achieved or revision limit reached"],
    ], [1.35, 3.4, 1.7])
    add_heading(doc, "State transition discipline", 2)
    add_body(doc, "Approval endpoints reject invalid transitions with HTTP 409 Conflict. This is expected protection: for example, content cannot be approved before an angle is selected, and video cannot be generated before slide approval. The UI must refresh its session after every transition to avoid submitting an action against stale state.")
    add_callout(doc, "Publication safeguard", "YouTube uploads use privacyStatus=private. Missing OAuth configuration does not block local completion; it prevents the external upload and reports the capability as unavailable.")
    page_break(doc)

    # Components
    add_title(doc, "6. Core components and integrations")
    add_table(doc, ["Component", "Implemented behavior", "Operational dependency"], [
        ["Planner", "Research routing and structured content intent", "OpenAI for live mode"],
        ["Research", "Web evidence with title, URL, publisher, score, and rank", "Search API key"],
        ["Angle generator", "Three distinct strategic angles", "OpenAI or deterministic demo"],
        ["Writer / critic", "Post generation and bounded revision loop", "OpenAI or deterministic demo"],
        ["Fact evaluator", "Claim extraction, passage comparison, human override", "Retrieved evidence"],
        ["Presentation", "Eight-slide PPTX based on supplied template", "Template.pptx"],
        ["Speaker notes", "Humanized narration without raw reference URLs", "Generated slide narrative"],
        ["Video composer", "Slides, avatar, transitions, timing, captions, MP4 checks", "HeyGen and FFmpeg"],
        ["YouTube", "OAuth upload with private visibility", "Client ID, secret, refresh token"],
        ["LangSmith", "Tracing and approved-content dataset export", "API key, project, dataset"],
    ], [1.35, 3.35, 1.85], compact=True)
    add_heading(doc, "Generated artifacts", 2)
    add_bullets(doc, [
        "Approved LinkedIn-style content stored with its session and evaluation record.",
        "Editable PowerPoint deck and slide-preview images.",
        "Humanized speaker notes used as narration input.",
        "Scene plan, SRT captions, composed MP4, and video validation metadata.",
        "Immutable JSON and HTML evaluation experiment reports.",
    ])
    add_heading(doc, "Brand and slide approach", 2)
    add_body(doc, "The presentation pipeline uses the supplied PowerPoint reference/template and a navy-blue and bright-gold visual system. Content fitting, line length, visual composition, and speaker-note generation are handled separately so on-slide copy stays concise while narration remains conversational.")
    page_break(doc)

    # Research and quality
    add_title(doc, "7. Research, grounding, and content quality")
    add_heading(doc, "Research decision", 2)
    add_body(doc, "Research is optional and controlled by request mode and planner output. When research is enabled, the system records retrieved sources and evaluates generated factual claims against their passages. When research is disabled, research-specific factual correctness is marked NOT_APPLICABLE rather than assigned an invented score.")
    add_heading(doc, "Recommended grounding model", 2)
    add_table(doc, ["Measure", "Question answered", "Current target"], [
        ["Retrieval Relevance", "Do retrieved sources relate to the requested topic?", ">= 0.50"],
        ["Evidence Coverage", "How many generated factual claims have adequate evidence?", ">= 0.90"],
        ["Faithfulness", "Does each claim accurately reflect its supporting evidence?", ">= 9/10"],
        ["Groundedness", "What is the overall evidence-backed accuracy of the content?", ">= 9/10"],
        ["Unsupported Claims", "Are any material claims unsupported?", "0 (hard gate)"],
        ["Factual Correctness", "Does researched content clear the factual gate?", ">= 7/10"],
    ], [1.45, 3.55, 1.5])
    add_callout(doc, "Metric rationale", "Precision@5 and Recall@5 were removed because unrestricted open-web retrieval does not provide a stable, bounded universe of all relevant documents. The project instead measures relevance, claim coverage, faithfulness, groundedness, and unsupported material claims.")
    add_heading(doc, "Claim review process", 2)
    add_numbered(doc, [
        "Extract factual claims from the generated content.",
        "Compare each claim with retrieved evidence passages and numerical details.",
        "Classify the claim as supported, partially supported, unsupported, contradicted, or needing review.",
        "Allow a human reviewer to override the automated classification with a reason.",
        "Re-evaluate the content and block approval while material unsupported claims remain.",
    ])
    page_break(doc)

    # Governance
    add_title(doc, "8. Human governance and controls")
    add_body(doc, "Human approval is a workflow requirement, not a cosmetic UI step. The backend checks current state before performing downstream actions, records every decision, and exposes reject/regenerate routes for each material artifact.")
    add_table(doc, ["Checkpoint", "Approval permits", "Rejection behavior"], [
        ["Angle", "Post generation", "Record reason and regenerate strategy/angles"],
        ["Content", "Library save and slide generation", "Return for revision or regeneration"],
        ["Claims", "Factual gate completion", "Human classification with rationale"],
        ["Slides", "Video generation", "Regenerate presentation with feedback"],
        ["Video", "Finalization and optional YouTube upload", "Regenerate composition/video"],
    ], [1.15, 2.65, 2.85])
    add_heading(doc, "Hard gates", 2)
    add_bullets(doc, [
        "Mandatory checkpoints cannot be bypassed.",
        "Material unsupported claims must equal zero.",
        "Contradicted claims must equal zero.",
        "Revision loops must remain within the configured maximum.",
        "Forbidden tool calls, false success states, secret exposure, and public YouTube uploads must equal zero.",
    ])
    add_heading(doc, "Security and privacy posture", 2)
    add_body(doc, "Secrets are read from local environment configuration and must not be committed or displayed in the UI. Provider failures should be reported without leaking credentials. External publication requires explicit user approval, and the YouTube adapter fixes the visibility to private.")
    add_callout(doc, "Operator responsibility", "Human approval confirms suitability and intent; it does not replace evidence review. Reviewers should inspect sources and resolve material claim warnings before approving researched content.", PALE_BLUE)
    page_break(doc)

    # Evaluation framework page 1
    add_title(doc, "9. Evaluation framework")
    add_body(doc, "The framework contains 24 metrics across five groups. Scores are stage-aware: unavailable evidence is represented as PENDING, NOT_EVALUATED, NOT_REQUIRED, or NOT_APPLICABLE rather than converted into a misleading zero or pass.")
    add_heading(doc, "Groups 1–3", 2)
    add_table(doc, ["Group", "Metric", "Pass rule / interpretation"], [
        ["Retrieval", "Retrieval Relevance", ">= 0.50; heuristic until human-calibrated"],
        ["Retrieval", "Evidence Coverage", ">= 0.90"],
        ["Factual", "Factual Correctness", ">= 7/10 when research is enabled; otherwise N/A"],
        ["Factual", "Groundedness", ">= 9/10"],
        ["Factual", "Faithfulness", ">= 9/10"],
        ["Factual", "Unsupported Claims", "0 material unsupported claims; hard gate"],
        ["Content", "Clarity", ">= 8/10"],
        ["Content", "Value", ">= 7/10"],
        ["Content", "Engagement", ">= 7/10 without rewarding clickbait"],
        ["Content", "Tone", ">= 8/10"],
        ["Content", "Originality", ">= 6/10"],
        ["Content", "Voice Match", ">= 7/10"],
        ["Content", "Audience Relevance", ">= 8/10"],
    ], [1.05, 1.7, 3.9], compact=True)
    add_heading(doc, "Evaluation semantics", 2)
    add_table(doc, ["Status", "Meaning"], [
        ["PASS", "Measured value meets the declared threshold"],
        ["REVISE / FAIL", "Measured value misses threshold or violates a gate"],
        ["PENDING", "Workflow has not reached the point where the metric can be evaluated"],
        ["NOT_EVALUATED", "Required labels or evidence are not yet available"],
        ["NOT_REQUIRED", "Metric does not apply to the selected route"],
        ["NOT_APPLICABLE", "Metric is structurally unsuitable for the scenario"],
    ], [1.35, 5.3], compact=True)
    page_break(doc)

    # Evaluation framework page 2
    add_title(doc, "9B. Remaining metrics")
    add_heading(doc, "Groups 4–5", 2)
    add_table(doc, ["Group", "Metric", "Evaluation approach"], [
        ["Human", "Human Edit %", "Deterministic difference between draft and approved final"],
        ["Human", "First Draft Approval", "Deterministic approval history"],
        ["Human", "Regeneration Count", "Deterministic workflow count"],
        ["Human", "Revision Count", "Deterministic workflow count"],
        ["Trajectory", "Correct Agent Path", "Expected versus actual LangGraph path"],
        ["Trajectory", "Tool Selection", "Required and forbidden tools by scenario"],
        ["Trajectory", "Tool Call Success", "Trace and provider-result evaluation"],
        ["Trajectory", "Step Efficiency", "Penalty for unnecessary calls"],
        ["Trajectory", "HITL Compliance", "Mandatory-checkpoint hard gate"],
        ["Trajectory", "Retry Behaviour", "Bounded retry and recovery checks"],
        ["Trajectory", "Goal Completion", "Valid requested terminal state reached"],
    ], [1.05, 1.75, 3.85], compact=True)
    add_heading(doc, "Evaluation assets", 2)
    add_table(doc, ["Asset", "Purpose"], [
        ["Golden Dataset", "40 cases: 20 happy path, 12 edge, 6 known failure, 2 adversarial"],
        ["Manifest", "Versioned split and dataset hash for reproducibility"],
        ["Policy", "Targets, content thresholds, hard gates, and version"],
        ["Session matrix", "Live evaluation state for one content session"],
        ["Experiment reports", "Immutable JSON and HTML summaries"],
        ["LangSmith", "Trace evidence and approved-content dataset export when configured"],
        ["pytest suite", "Workflow, metrics, artifacts, gates, and regression checks"],
    ], [1.65, 5.0], compact=True)
    page_break(doc)

    # Current evaluation status
    add_title(doc, "10. Current evaluation status")
    add_table(doc, ["Evidence", "Result", "Interpretation"], [
        ["Automated test suite", "31/31 passed", "Implemented behaviors are regression-protected"],
        ["Validation experiment", "8/8 passed", "Controlled workflow cases reached valid outcomes"],
        ["Completion rate", "100%", "Exceeds the 95% workflow target in deterministic mode"],
        ["Provider cost", "$0", "Validation used deterministic provider substitutions"],
        ["Faithfulness", "NOT_EVALUATED", "Human claim/evidence mappings are still required"],
        ["Evidence Coverage", "NOT_EVALUATED", "Controlled cases do not yet contain adjudicated claim coverage"],
        ["Release ready", "False", "Correctly blocked by the missing faithfulness evidence"],
    ], [1.65, 1.35, 3.65])
    add_heading(doc, "Readiness assessment", 2)
    add_table(doc, ["Dimension", "Score", "Rationale"], [
        ["Evaluation architecture", "9/10", "Versioned policy, reports, hard gates, stage-aware states"],
        ["Golden-dataset design", "8/10", "Good scenario distribution and fixed splits"],
        ["Workflow trajectory", "9/10", "Strong route, tool, retry, and HITL checks"],
        ["Content quality", "7/10", "Useful rubric; human calibration remains incomplete"],
        ["Grounding evidence", "5/10", "Evaluator exists, but validation labels are incomplete"],
        ["Reproducibility", "8.5/10", "Hashes, versions, zero-cost runner, immutable artifacts"],
        ["Overall", "7.5/10", "Evaluation-capable, not yet evidence-complete for release"],
    ], [1.75, 0.85, 4.05])
    add_callout(doc, "Do not overclaim", "The 8/8 result demonstrates deterministic workflow conformance. It does not establish live search quality, live LLM faithfulness, or production latency and cost.", PALE_GOLD)
    add_title(doc, "10B. Release evidence", size=30)
    add_bullets(doc, [
        "Human-adjudicate approximately 30–50 representative claim/evidence pairs.",
        "Demonstrate at least 85% evaluator agreement and zero material false passes.",
        "Add 10–20 approved Ekmind posts as a stable voice-reference corpus.",
        "Calibrate Voice Match and content-quality judgments against human ratings.",
        "Run the validation set with live providers, then freeze the evaluator version.",
        "Run and publish the hidden-test report with trace, latency, token, and cost evidence.",
    ])
    add_table(doc, ["Gate", "Required evidence"], [
        ["Calibration", "Human agreement >=85%; zero material false passes"],
        ["Voice", "Approved-post reference set and human-rated comparison"],
        ["Live validation", "Trace, latency, token, cost, and provider outcomes"],
        ["Final release", "Frozen evaluator; untouched hidden-test split passes"],
    ], [1.5, 5.15])
    add_callout(doc, "Decision rule", "If any material false pass occurs, or the hidden-test run requires prompt or evaluator changes, the release remains blocked and a new versioned validation cycle is required.", PALE_BLUE)
    page_break(doc)

    # Setup and testing
    add_title(doc, "11. Setup, operation, and testing")
    add_heading(doc, "Local startup", 2)
    add_numbered(doc, [
        r"Open PowerShell in ekmind-content-studio.",
        r"Start the API: .venv\Scripts\python -m uvicorn backend.main:app --reload",
        r"Open a second PowerShell window.",
        r"Start the UI: .venv\Scripts\python -m streamlit run frontend/streamlit_app.py",
        r"Open http://localhost:8501 and verify http://127.0.0.1:8000/health.",
    ])
    add_heading(doc, "Functional acceptance test", 2)
    add_table(doc, ["Step", "Expected observation"], [
        ["Create session", "Topic accepted and three distinct angles returned"],
        ["Approve/reject angle", "Valid transition succeeds; reason is recorded on rejection"],
        ["Generate content", "Draft, critique, revisions, sources, and metrics appear"],
        ["Review claims", "Claim classifications can be overridden and re-evaluated"],
        ["Approve content", "Approved item becomes available in the library"],
        ["Generate slides", "Eight previews, PPTX, fitting content, and speaker notes appear"],
        ["Approve slides", "Video generation becomes available"],
        ["Generate video", "MP4, audio stream, captions, and validation metadata exist"],
        ["Approve video", "Local workflow completes; private upload occurs only if configured"],
        ["Open Evaluations", "24 metrics shown; Precision@5 and Recall@5 absent"],
    ], [1.65, 5.0], compact=True)
    add_heading(doc, "Automated checks", 2)
    add_body(doc, r"Run .venv\Scripts\python -m pytest -q. The current expected result is 31 passed. Generate a new zero-cost validation report with a unique experiment ID using:")
    add_callout(doc, "Command", r".venv\Scripts\python -m backend.evaluations.cli --split validation --experiment <unique_id>", PALE_BLUE)
    page_break(doc)

    # Risks and roadmap
    add_title(doc, "12. Risks, limitations, and roadmap")
    add_heading(doc, "Current risks and mitigations", 2)
    add_table(doc, ["Risk", "Impact", "Mitigation / next action"], [
        ["Uncalibrated faithfulness", "Material claims could be misclassified", "Human-label claim/evidence pairs; measure agreement and false passes"],
        ["Heuristic voice score", "Content may not consistently sound like Ekmind", "Add approved posts and calibrate against human ratings"],
        ["Mock validation", "Live cost, latency, and provider failure modes remain uncertain", "Run controlled live-provider validation and capture LangSmith traces"],
        ["SQLite baseline", "Limited concurrency and production resilience", "Move to PostgreSQL and introduce managed migrations"],
        ["External provider credits", "HeyGen generation may fail despite valid credentials", "Preflight credit checks and actionable error reporting"],
        ["Template variability", "PowerPoint rendering may differ across Office versions", "Validate on target PowerPoint environment and keep overflow tests"],
        ["OAuth lifecycle", "YouTube refresh tokens may expire or be revoked", "Health check and guided re-authorization"],
    ], [1.7, 2.1, 2.85], compact=True)
    add_heading(doc, "Prioritized roadmap", 2)
    add_table(doc, ["Priority", "Improvement", "Measured success"], [
        ["P0", "Complete faithfulness calibration", ">=85% agreement; zero material false passes"],
        ["P0", "Run live validation", "Published traces, latency, tokens, cost, and provider outcomes"],
        ["P1", "Build voice-reference set", "10–20 approved posts; Voice Match agreement validated"],
        ["P1", "Freeze evaluator version", "Reproducible validation and hidden-test results"],
        ["P1", "Production persistence", "PostgreSQL migration and workspace isolation tests"],
        ["P2", "Operational resilience", "Provider retries, quota checks, and recovery dashboards"],
        ["P2", "Publishing expansion", "Only after governance review; LinkedIn remains out of scope"],
    ], [0.75, 2.7, 3.2], compact=True)
    add_callout(doc, "Recommended release condition", "Do not declare evaluation-ready for production until faithfulness is measured on human-adjudicated data and the hidden-test split passes without changing prompts, thresholds, or evaluator logic.")
    page_break(doc)

    # Appendix
    add_title(doc, "Appendix A. API and configuration reference")
    add_heading(doc, "Principal API endpoints", 2)
    add_table(doc, ["Method", "Endpoint", "Purpose"], [
        ["GET", "/health", "Capability and configuration health"],
        ["POST", "/api/sessions", "Create strategy session"],
        ["GET", "/api/sessions/{id}", "Retrieve current session"],
        ["POST", "/api/sessions/{id}/approve-angle", "Approve selected angle"],
        ["POST", "/api/sessions/{id}/approve-content", "Approve edited/final content"],
        ["POST", "/api/sessions/{id}/review-claim", "Record human claim judgment"],
        ["POST", "/api/sessions/{id}/reevaluate-content", "Re-run claim/evidence assessment"],
        ["POST", "/api/sessions/{id}/slides", "Generate presentation"],
        ["POST", "/api/sessions/{id}/approve-slides", "Approve slides"],
        ["POST", "/api/sessions/{id}/video", "Generate video"],
        ["POST", "/api/sessions/{id}/approve-video", "Approve video and optional upload"],
        ["GET", "/api/evaluations/experiments", "List experiment reports"],
        ["POST", "/api/sessions/{id}/export-langsmith", "Export approved example"],
    ], [0.72, 3.25, 2.68], compact=True)
    add_heading(doc, "Configuration groups", 2)
    add_table(doc, ["Area", "Variables / configuration"], [
        ["LLM", "OPENAI_API_KEY and model selection"],
        ["Research", "SEARCH_API_KEY / configured search provider"],
        ["LangSmith", "LANGSMITH_API_KEY, project, dataset, tracing flags"],
        ["Database", "DATABASE_URL; SQLite default, PostgreSQL optional"],
        ["HeyGen", "API key, avatar ID, and voice ID"],
        ["YouTube", "YOUTUBE_CLIENT_ID, YOUTUBE_CLIENT_SECRET, YOUTUBE_REFRESH_TOKEN"],
        ["Brand", "Template.pptx and brand_config.json"],
    ], [1.25, 5.4], compact=True)
    add_title(doc, "Appendix B. Artifacts")
    add_bullets(doc, [
        "README.md — implementation and startup overview.",
        "evaluation_data/evaluation_policy.json — metric targets and hard gates.",
        "evaluation_data/golden_manifest.json — versioned dataset split and hashes.",
        "storage/experiments/free_validation_v2.json and .html — latest controlled report.",
        "Ekmind_AI_Content_Studio_Golden_Dataset.xlsx — authoritative 40-case dataset.",
        "Template.pptx — source presentation format; original is not modified.",
    ])
    add_callout(doc, "Document baseline", "Implementation evidence reviewed on 4 September 2026. Future changes to prompts, providers, policies, templates, or evaluators should trigger a new versioned experiment and a document update.", PALE_BLUE)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
